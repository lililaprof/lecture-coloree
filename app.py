import streamlit as st
from PIL import Image, ImageOps, ImageFilter
import pytesseract
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# Configuration de base
st.set_page_config(page_title="Lecture Colorée CP", page_icon="📚", layout="wide")

# Google Analytics
import streamlit.components.v1 as components
components.html("""
<!-- Google tag (gtag.js) -->
<script async src="https://www.googletagmanager.com/gtag/js?id=G-GKY6GERLTX"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){dataLayer.push(arguments);}
  gtag('js', new Date());
  gtag('config', 'G-GKY6GERLTX');
</script>
""", height=0)

# Définitions globales
sons_complexes = [
    'ouil', 'euil', 'aille', 'eille', 'ille', 'ouille',
    'ain', 'aim', 'ein', 'eim', 'oin', 'ien', 'eau', 'oeu',
    'ch', 'ph', 'gn', 'ill', 'ail', 'eil', 'ou', 'au', 'eu', 'oi', 'oy',
    'ai', 'ei'
]

sons_nasals = ['an', 'am', 'en', 'em', 'on', 'om', 'in', 'im', 'un', 'um', 'yn', 'ym']
voyelles = 'aeiouyàâäéèêëïîôùûüÿæœAEIOUYÀÂÄÉÈÊËÏÎÔÙÛÜŸÆŒ'
lettres_muettes_fin = ['s', 't', 'd', 'p', 'x', 'z']

LISTES_MANUELS = {
    'Ma liste perso': [],
    'Taoki': ['est', 'et', 'un', 'une', 'le', 'la', 'les', 'de', 'il', 'elle', 'dans', 'sur', 'avec'],
    'Pilotis': ['le', 'la', 'les', 'un', 'une', 'des', 'il', 'elle', 'est', 'dans', 'sur', 'avec', 'pour'],
    'Léo et Léa': ['le', 'la', 'l', 'un', 'une', 'et', 'est', 'il', 'elle', 'je', 'tu', 'de', 'du'],
    'Base commune': ['est', 'et', 'un', 'une', 'le', 'la', 'les', 'de', 'du', 'des', 'dans', 'sur', 
                     'avec', 'pour', 'par', 'il', 'elle', 'ils', 'elles', 'ont', 'sont', 'a', 'à', 
                     'au', 'aux', 'ce', 'cette', 'ces', 'mon', 'ma', 'mes', 'ton', 'ta', 'tes', 'son', 'sa', 'ses']
}

POLICES = ['Arial', 'Comic Sans MS', 'OpenDyslexic', 'Quicksand Book', 'Belle Allure', 'Helvetica']

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return RGBColor(*tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4)))

def pretraiter_image(image, qualite_ocr):
    img = image.copy()
    img = ImageOps.grayscale(img)
    if qualite_ocr == "Maximale":
        img = img.filter(ImageFilter.MedianFilter(size=3))
        img = img.point(lambda p: 255 if p > 180 else 0)
    return img

def nettoyer_texte_ocr(texte):
    texte = re.sub(r'\s+[•|]\s+', ' ', texte)
    texte = re.sub(r'\s+\d+\s+(?=[a-z]{1,2}\s)', ' ', texte)
    texte = re.sub(r'\s+[a-zA-Z]{1,2}\s+(?=[a-z])', ' ', texte)
    texte = re.sub(r'\s+', ' ', texte)
    lignes = texte.split('\n')
    lignes_propres = [l for l in lignes if l.strip() and not l.strip().startswith(('•', '|', '-', '=', '_'))]
    return '\n'.join(lignes_propres).strip()

def detecter_lettre_muette(mot, position):
    mot_lower = mot.lower()
    if position == len(mot) - 1:
        if mot[position].lower() in lettres_muettes_fin:
            return True
    if len(mot) >= 3 and mot_lower.endswith('ent'):
        if position >= len(mot) - 3:
            return True
    if position == 0 and mot[position].lower() == 'h':
        return True
    return False

def extraire_mot_complet(texte, position):
    debut = position
    while debut > 0 and texte[debut - 1].isalpha():
        debut -= 1
    fin = position
    while fin < len(texte) and texte[fin].isalpha():
        fin += 1
    return texte[debut:fin], debut, fin

def est_son_nasal_valide(texte, position, son):
    pos_apres = position + len(son)
    if pos_apres >= len(texte):
        return True
    char_apres = texte[pos_apres].lower()
    if char_apres in voyelles:
        return False
    if son[-1] == 'n' and char_apres == 'n':
        return False
    if son[-1] == 'm' and char_apres == 'm':
        return False
    return True

def remplacer_separateurs(texte):
    resultat = ""
    i = 0
    while i < len(texte):
        if texte[i] == '.':
            if i + 1 < len(texte):
                reste = texte[i+1:].lstrip()
                if reste and (reste[0].isupper() or reste[0] == '\n'):
                    resultat += '.'
                else:
                    resultat += ' •'
            else:
                resultat += '.'
            i += 1
        else:
            resultat += texte[i]
            i += 1
    return resultat

def ajouter_espaces_entre_mots(texte):
    resultat = ""
    for i, char in enumerate(texte):
        if char == ' ':
            if i > 0 and texte[i-1] != ' ':
                resultat += '  '
            elif i == 0:
                resultat += '  '
        else:
            resultat += char
    return resultat

def mettre_majuscules_phrases(texte):
    resultat = []
    debut_phrase = True
    for char in texte:
        if debut_phrase and char.isalpha():
            resultat.append(char.upper())
            debut_phrase = False
        else:
            resultat.append(char)
            if char in '.!?':
                debut_phrase = True
    return ''.join(resultat)

def colorier_texte(texte, mots_outils, couleurs_config, activer_muettes=True):
    resultat_word = []
    tous_mots = set(mots_outils + [m.lower() for m in mots_outils] + [m.upper() for m in mots_outils])
    i = 0
    while i < len(texte):
        char = texte[i]
        if not char.isalpha():
            resultat_word.append((char, None))
            i += 1
            continue
        mot_complet, debut_mot, fin_mot = extraire_mot_complet(texte, i)
        position_dans_mot = i - debut_mot
        if mot_complet in tous_mots:
            for c in mot_complet:
                resultat_word.append((c, 'mots_outils'))
            i = fin_mot
            continue
        if activer_muettes and detecter_lettre_muette(mot_complet, position_dans_mot):
            resultat_word.append((char, 'muettes'))
            i += 1
            continue
        trouve = False
        for son in sons_complexes:
            if texte[i:i+len(son)].lower() == son:
                for c in texte[i:i+len(son)]:
                    resultat_word.append((c, 'graphemes'))
                i += len(son)
                trouve = True
                break
        if not trouve:
            for son in sons_nasals:
                if texte[i:i+len(son)].lower() == son:
                    if est_son_nasal_valide(texte, i, son):
                        for c in texte[i:i+len(son)]:
                            resultat_word.append((c, 'graphemes'))
                        i += len(son)
                        trouve = True
                        break
        if not trouve:
            resultat_word.append((char, 'voyelles' if char.lower() in voyelles else 'consonnes'))
            i += 1
    return resultat_word

def colorier_texte_simple_options(texte, mots_outils, col_graphemes, col_mots_outils, 
                                   activer_graphemes=False, activer_mots_outils=False):
    resultat_word = []
    tous_mots = set(mots_outils + [m.lower() for m in mots_outils] + [m.upper() for m in mots_outils])
    i = 0
    while i < len(texte):
        char = texte[i]
        if not char.isalpha():
            resultat_word.append((char, None))
            i += 1
            continue
        mot_complet, debut_mot, fin_mot = extraire_mot_complet(texte, i)
        if activer_mots_outils and mot_complet in tous_mots:
            for c in mot_complet:
                resultat_word.append((c, 'mots_outils'))
            i = fin_mot
            continue
        trouve = False
        if activer_graphemes:
            for son in sons_complexes:
                if texte[i:i+len(son)].lower() == son:
                    for c in texte[i:i+len(son)]:
                        resultat_word.append((c, 'graphemes'))
                    i += len(son)
                    trouve = True
                    break
            if not trouve:
                for son in sons_nasals:
                    if texte[i:i+len(son)].lower() == son:
                        if est_son_nasal_valide(texte, i, son):
                            for c in texte[i:i+len(son)]:
                                resultat_word.append((c, 'graphemes'))
                            i += len(son)
                            trouve = True
                            break
        if not trouve:
            resultat_word.append((char, None))
            i += 1
    return resultat_word

def colorier_graphemes_cibles(texte, graphemes_cibles, couleur_cible):
    resultat_word = []
    graphemes_lower = [g.lower() for g in graphemes_cibles]
    i = 0
    while i < len(texte):
        trouve = False
        for grapheme in graphemes_lower:
            if texte[i:i+len(grapheme)].lower() == grapheme:
                for c in texte[i:i+len(grapheme)]:
                    resultat_word.append((c, 'cible'))
                i += len(grapheme)
                trouve = True
                break
        if not trouve:
            resultat_word.append((texte[i], 'black'))
            i += 1
    return resultat_word

def creer_word(texte_traite, police, couleurs_config, casse, taille_pt=25):
    doc = Document()
    couleurs_rgb = {k: hex_to_rgb(v) for k, v in couleurs_config.items()}
    titre_para = doc.add_heading(casse.upper(), level=1)
    titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for char, couleur in texte_traite:
        run = para.add_run(char)
        run.font.size = Pt(taille_pt)
        run.font.name = police
        if couleur and couleur in couleurs_rgb:
            run.font.color.rgb = couleurs_rgb[couleur]
    return doc

def generer_preview_html(texte_traite, couleurs_config, police):
    mapping = {
        'voyelles': couleurs_config.get('voyelles', '#FF0000'),
        'consonnes': couleurs_config.get('consonnes', '#0000FF'),
        'graphemes': couleurs_config.get('graphemes', '#008000'),
        'muettes': couleurs_config.get('muettes', '#808080'),
        'mots_outils': couleurs_config.get('mots_outils', '#8B4513'),
        'cible': couleurs_config.get('cible', '#069494'),
        'black': '#000000',
        None: '#000000'
    }
    html = f"<div style='font-family:{police}; font-size:18px; line-height:1.8; padding:15px; background:#f9f9f9; border-radius:10px;'>"
    for char, couleur in texte_traite:
        color = mapping.get(couleur, '#000000')
        safe_char = char.replace(' ', '&nbsp;').replace('\n', '<br/>')
        html += f"<span style='color:{color};'>{safe_char}</span>"
    html += "</div>"
    return html

# Interface
st.title("📚 Lecture Colorée pour CP")
st.markdown("**Application d'adaptation de textes pour enfants dys et TSA**")
st.markdown("*Pour les enseignants et les parents*")

st.info("""
📖 **Comment ça marche ?**
1. Uploadez une ou plusieurs photos OU tapez/collez votre texte
2. Vérifiez et corrigez le texte extrait si besoin
3. Choisissez les documents à générer
4. Téléchargez !

🎨 **Code couleur :** 🔴 Voyelles • 🔵 Consonnes • 🟢 Graphèmes complexes • ⚫ Lettres muettes • 🟤 Mots-outils
""")

with st.expander("ℹ️ En savoir plus"):
    st.markdown("""
    ### Nouveautés
    - ✅ Upload multiple d'images
    - ✅ Zone de correction du texte
    - ✅ Nettoyage automatique
    - ✅ Détection "ent" final améliorée
    
    *Application gratuite et open source* 💚
    """)

st.markdown("---")

# Sidebar
with st.sidebar:
    st.header("⚙️ Paramètres")
    police = st.selectbox("📝 Police", POLICES, index=0)
    taille_police = st.slider("📏 Taille (pt)", 12, 40, 25, 1)
    casse = st.radio("📝 Casse", ['Minuscules', 'Majuscules'], horizontal=True)
    st.markdown("---")
    qualite_ocr = st.select_slider("🔍 Qualité OCR", ["Standard", "Bonne", "Maximale"], "Bonne")

# Zone input
st.header("📥 Votre texte")
tab1, tab2 = st.tabs(["📤 Upload d'image(s)", "✍️ Saisie manuelle"])

with tab1:
    uploaded_files = st.file_uploader("Image(s) - Plusieurs possibles !", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True)
    if uploaded_files:
        st.info(f"✅ {len(uploaded_files)} image(s)")
        cols = st.columns(min(len(uploaded_files), 4))
        for idx, file in enumerate(uploaded_files):
            with cols[idx % 4]:
                st.image(Image.open(file), caption=f"Image {idx+1}", use_column_width=True)

with tab2:
    texte_saisi = st.text_area("Tapez ou collez votre texte", height=200)

# Extraction OCR
if uploaded_files:
    if st.button("🔍 Extraire le texte des images", type="primary"):
        with st.spinner("📖 Extraction..."):
            textes = []
            for file in uploaded_files:
                img = Image.open(file)
                img_pre = pretraiter_image(img, qualite_ocr)
                config = '--psm 6 -l fra' if qualite_ocr != "Maximale" else '--psm 3 -l fra'
                texte = pytesseract.image_to_string(img_pre, config=config)
                textes.append(nettoyer_texte_ocr(texte))
            st.session_state['texte_extrait'] = "\n\n".join(textes)
            st.success(f"✅ Texte extrait de {len(uploaded_files)} image(s) !")

# Zone éditable
if 'texte_extrait' in st.session_state or texte_saisi:
    st.markdown("---")
    st.header("✏️ Texte à traiter")
    texte_init = st.session_state.get('texte_extrait', texte_saisi)
    texte_editable = st.text_area("Vérifiez et corrigez", value=texte_init, height=300)
    
    if texte_editable:
        st.markdown("---")
        st.header("📄 Documents à générer")
        
        # Option 1
        st.markdown("### 🎨 Texte coloré")
        creer_texte_colore = st.toggle("Activer", key="colore", value=True)
        if creer_texte_colore:
            st.info("📖 Code couleur complet")
            col1, col2, col3 = st.columns(3)
            with col1:
                col_voy = st.color_picker("🔴 Voyelles", "#FF0000")
                col_cons = st.color_picker("🔵 Consonnes", "#0000FF")
            with col2:
                col_graph = st.color_picker("🟢 Graphèmes", "#008000")
                col_muet = st.color_picker("⚫ Muettes", "#808080")
            with col3:
                col_mots = st.color_picker("🟤 Mots-outils", "#8B4513")
                activer_muettes = st.checkbox("Détecter muettes", True)
            manuel = st.selectbox("Liste mots-outils", list(LISTES_MANUELS.keys()))
            mots_base = LISTES_MANUELS[manuel].copy()
            if manuel == 'Ma liste perso':
                perso = st.text_area("Vos mots", "", height=60)
                if perso:
                    mots_base.extend([m.strip() for m in perso.split(',') if m.strip()])
            couleurs_config = {'voyelles': col_voy, 'consonnes': col_cons, 'graphemes': col_graph, 'muettes': col_muet, 'mots_outils': col_mots}
        
        st.markdown("---")
        
        # Option 2
        st.markdown("### 📃 Texte simple")
        creer_simple = st.toggle("Activer", key="simple")
        if creer_simple:
            st.info("📝 Noir/blanc + options")
            col1, col2 = st.columns(2)
            with col1:
                simple_graph = st.checkbox("Graphèmes")
                col_graph_simple = st.color_picker("Couleur", "#008000", key="cgs") if simple_graph else "#008000"
            with col2:
                simple_mots = st.checkbox("Mots-outils")
                if simple_mots:
                    col_mots_simple = st.color_picker("Couleur", "#8B4513", key="cms")
                    manuel_s = st.selectbox("Liste", list(LISTES_MANUELS.keys()), key="ms")
                    mots_s = LISTES_MANUELS[manuel_s].copy()
                else:
                    col_mots_simple = "#8B4513"
                    mots_s = []
        
        st.markdown("---")
        
        # Option 3
        st.markdown("### 🎯 Graphèmes ciblés")
        creer_cible = st.toggle("Activer", key="cible")
        if creer_cible:
            st.info("🎯 Son spécifique en couleur")
            col1, col2 = st.columns([2, 1])
            with col1:
                graph_input = st.text_input("Graphèmes", placeholder="ou, ch, ain")
            with col2:
                col_cible = st.color_picker("Couleur", "#069494")
        
        st.markdown("---")
        
        # Génération
        if st.button("🚀 GÉNÉRER", type="primary", use_container_width=True):
            if not (creer_simple or creer_texte_colore or creer_cible):
                st.warning("⚠️ Activez au moins un document !")
            else:
                with st.spinner("⏳ Génération..."):
                    try:
                        texte = remplacer_separateurs(texte_editable)
                        texte = ajouter_espaces_entre_mots(texte)
                        texte = texte.lower() if casse == 'Minuscules' else texte.upper()
                        if casse == 'Minuscules':
                            texte = mettre_majuscules_phrases(texte)
                        
                        if creer_simple:
                            st.info("📄 Texte simple...")
                            t_simple = colorier_texte_simple_options(texte, mots_s, col_graph_simple, col_mots_simple, simple_graph, simple_mots)
                            st.subheader("👁️ Aperçu")
                            st.markdown(generer_preview_html(t_simple, {'graphemes': col_graph_simple, 'mots_outils': col_mots_simple}, police), unsafe_allow_html=True)
                            doc = creer_word(t_simple, police, {'graphemes': col_graph_simple, 'mots_outils': col_mots_simple}, casse, taille_police)
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            st.download_button("📥 Télécharger", buf, f"simple_{casse.lower()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        
                        if creer_texte_colore:
                            st.info("📄 Texte coloré...")
                            t_col = colorier_texte(texte, mots_base, couleurs_config, activer_muettes)
                            st.subheader("👁️ Aperçu")
                            st.markdown(generer_preview_html(t_col, couleurs_config, police), unsafe_allow_html=True)
                            doc = creer_word(t_col, police, couleurs_config, casse, taille_police)
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            st.download_button("📥 Télécharger", buf, f"colore_{casse.lower()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        
                        if creer_cible and graph_input:
                            st.info("📄 Graphèmes ciblés...")
                            graphs = [g.strip() for g in graph_input.split(',') if g.strip()]
                            if graphs:
                                t_cib = colorier_graphemes_cibles(texte, graphs, col_cible)
                                st.subheader("👁️ Aperçu")
                                st.markdown(generer_preview_html(t_cib, {'cible': col_cible, 'black': '#000000'}, police), unsafe_allow_html=True)
                                doc = creer_word(t_cib, police, {'cible': col_cible, 'black': '#000000'}, casse, taille_police)
                                buf = io.BytesIO()
                                doc.save(buf)
                                buf.seek(0)
                                st.download_button("📥 Télécharger", buf, f"cible_{casse.lower()}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        
                        st.success("🎉 Terminé !")
                    except Exception as e:
                        st.error(f"❌ Erreur : {str(e)}")

st.markdown("---")
st.markdown("*Créé avec ❤️ - Projet open source*")
